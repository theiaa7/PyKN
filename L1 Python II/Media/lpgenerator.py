from docx import Document
from docx.shared import Inches

document = Document()

document.add_heading('Lesson 01: Class and Method', 0)

p = document.add_paragraph('Learning python class and method for implementation game development.')
p.add_run('bold').bold = True
p.add_run('italic.').italic = True

records = (
    # (10, '101', game_01.Player.__doc__, ''),
    (5, '422', 'Eggs', ''),
    (5, '631', 'Spam, spam, eggs, and spam', ''),
    (10, '631', 'Spam, spam, eggs, and spam', ''),
    (5, '631', 'Spam, spam, eggs, and spam', ''),
)

table = document.add_table(rows=1, cols=4)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Timing'
hdr_cells[1].text = 'Lesson Description'
hdr_cells[2].text = 'Additional Material'
hdr_cells[3].text = 'Comments'

for qty, id, desc, comm in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc
    row_cells[3].text = comm

document.add_page_break()

document.save('demo.docx')